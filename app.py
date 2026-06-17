"""Facility Assessment Snapshot -- INFINITE by MEDELITE
Bonus features: Word doc export, metric cards/charts, advanced error handling.
"""

import streamlit as st
import requests
import pandas as pd
from io import BytesIO

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                 Paragraph, Spacer, HRFlowable)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT

st.set_page_config(
    page_title="Facility Assessment Snapshot — INFINITE by MEDELITE",
    page_icon="🏥", layout="wide"
)

st.markdown("""
<style>
.infinite-banner {
    background-color:#003366;color:white;padding:0.75rem 1rem;
    text-align:center;font-weight:bold;font-size:1.15rem;
    border-radius:6px;margin-bottom:1.2rem;letter-spacing:0.03em;
}
.metric-card {
    background:#f8fafc;border:1px solid #e2e8f0;border-radius:8px;
    padding:1rem;text-align:center;
}
.metric-label { font-size:0.75rem;color:#64748b;font-weight:600;
    text-transform:uppercase;letter-spacing:0.05em;margin-bottom:0.25rem; }
.metric-value { font-size:1.6rem;font-weight:700;color:#003366; }
.metric-sub   { font-size:0.7rem;color:#94a3b8;margin-top:0.1rem; }
.star-1 { color:#ef4444; } .star-2 { color:#f97316; }
.star-3 { color:#eab308; } .star-4 { color:#22c55e; } .star-5 { color:#16a34a; }
.hosp-table th { background:#003366;color:white;padding:6px 10px;font-size:0.8rem; }
.hosp-table td { padding:5px 10px;font-size:0.82rem;border-bottom:1px solid #f1f5f9; }
.hosp-table tr:nth-child(even) td { background:#fffbeb; }
.section-header { font-size:0.7rem;font-weight:700;color:#64748b;
    text-transform:uppercase;letter-spacing:0.08em;
    border-bottom:2px solid #003366;padding-bottom:4px;margin-bottom:8px; }
</style>
<div class="infinite-banner">INFINITE — Managed by MEDELITE</div>
""", unsafe_allow_html=True)

# ── Dataset IDs ────────────────────────────────────────────────────────────────
PROVIDER_DS  = "4pq5-n9py"
STATE_AVG_DS = "xcdc-v8bm"
API_BASE     = "https://data.cms.gov/provider-data/api/1/datastore/query"

# ── Confirmed column names from xcdc-v8bm ─────────────────────────────────────
COL_STR_HOSP = "percentage_of_short_stay_residents_who_were_rehospitalized__1d02"
COL_STR_ED   = "percentage_of_short_stay_residents_who_had_an_outpatient_em_d911"
COL_LT_HOSP  = "number_of_hospitalizations_per_1000_longstay_resident_days"
COL_LT_ED    = "number_of_outpatient_emergency_department_visits_per_1000_l_de9d"

HOSP_FIELDS = [
    ("Short Term Hospitalization",            COL_STR_HOSP, "facility", True),
    ("STR National Avg. for Hospitalization", COL_STR_HOSP, "national", True),
    ("STR State Avg. for Hospitalization",    COL_STR_HOSP, "state",    True),
    ("STR ED Visit",                          COL_STR_ED,   "facility", True),
    ("STR ED Visits National Avg.",           COL_STR_ED,   "national", True),
    ("STR ED Visits State Avg.",              COL_STR_ED,   "state",    True),
    ("LT Hospitalization",                    COL_LT_HOSP,  "facility", False),
    ("LT National Avg. for Hospitalization",  COL_LT_HOSP,  "national", False),
    ("LT State Avg. for Hospitalization",     COL_LT_HOSP,  "state",    False),
    ("ED Visit",                              COL_LT_ED,    "facility", False),
    ("LT ED Visits National Avg.",            COL_LT_ED,    "national", False),
    ("LT ED Visits State Avg.",               COL_LT_ED,    "state",    False),
]


# ── python-docx helpers ────────────────────────────────────────────────────────

def _docx_set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def _docx_set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def _docx_add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def _docx_add_row(table, label, value, highlight=False):
    row = table.add_row()
    lc = row.cells[0]
    lc.text = ""
    p = lc.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _docx_set_cell_bg(lc, 'FFFF00' if highlight else 'F7F7F7')
    _docx_set_cell_borders(lc)
    vc = row.cells[1]
    vc.text = ""
    p2 = vc.paragraphs[0]
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(2)
    run2 = p2.add_run(str(value) if value else 'N/A')
    run2.italic = True
    run2.font.size = Pt(9)
    run2.font.name = 'Arial'
    _docx_set_cell_bg(vc, 'FFFF00' if highlight else 'FFFFFF')
    _docx_set_cell_borders(vc)


# ── API helpers ────────────────────────────────────────────────────────────────

def _post(dataset, payload, timeout=20):
    r = requests.post(f"{API_BASE}/{dataset}/0", json=payload, timeout=timeout)
    return r.json().get("results", []) if r.status_code == 200 else []

def _get(dataset, params, timeout=20):
    r = requests.get(f"{API_BASE}/{dataset}/0", params=params, timeout=timeout)
    return r.json().get("results", []) if r.status_code == 200 else []

def _fetch_by_ccn(dataset, ccn, limit=5):
    ccn = ccn.strip()
    for attempt in ["post", "get"]:
        try:
            rows = (_post(dataset, {"conditions":[{"property":"cms_certification_number_ccn",
                                                   "value":ccn,"operator":"="}],"limit":limit})
                    if attempt == "post" else
                    _get(dataset, {"conditions[0][property]":"cms_certification_number_ccn",
                                   "conditions[0][value]":ccn,
                                   "conditions[0][operator]":"=","limit":limit}))
            matched = [r for r in rows
                       if str(r.get("cms_certification_number_ccn","")).strip() == ccn]
            if matched:
                return matched
        except Exception:
            pass
    # Page scan fallback
    try:
        offset = 0
        while offset < 20000:
            rows = _get(dataset, {"limit":500,"offset":offset})
            if not rows:
                break
            for row in rows:
                if str(row.get("cms_certification_number_ccn","")).strip() == ccn:
                    return [row]
            offset += 500
    except Exception:
        pass
    return []

def fetch_provider_data(ccn):
    rows = _fetch_by_ccn(PROVIDER_DS, ccn, limit=1)
    return rows[0] if rows else {}

@st.cache_data(ttl=86400, show_spinner=False)
def fetch_averages(state):
    out = {"national":{}, "state":{}}
    for scope in [state.upper(), "NATION"]:
        try:
            rows = _post(STATE_AVG_DS, {"conditions":[{"property":"state_or_nation",
                                                        "value":scope,"operator":"="}],"limit":1})
            if not rows:
                rows = _get(STATE_AVG_DS, {"conditions[0][property]":"state_or_nation",
                                            "conditions[0][value]":scope,
                                            "conditions[0][operator]":"=","limit":1})
            if rows:
                bucket = "national" if scope == "NATION" else "state"
                out[bucket] = {k: str(v) for k, v in rows[0].items()}
        except Exception:
            pass
    return out

def resolve_hosp(col, source, averages, is_pct):
    if source == "facility":
        return "Not Reported"
    val = averages.get(source, {}).get(col, "")
    if not val or val.lower() in ("nan","none",""):
        return "N/A"
    try:
        f = float(val)
        return f"{round(f,1)}%" if is_pct else str(round(f,2))
    except Exception:
        return val

def _v(d, *keys):
    for k in keys:
        val = str(d.get(k,"")).strip()
        if val and val.lower() not in ("nan","none",""):
            return val
    return "N/A"

def validate_ccn(ccn):
    ccn = ccn.strip()
    if not ccn:
        return None, "Please enter a CCN."
    if not ccn.isdigit():
        return None, f"CCN must contain only digits. Got: '{ccn}'"
    if len(ccn) != 6:
        return None, f"CCN must be exactly 6 digits. Got {len(ccn)} digits: '{ccn}'"
    return ccn, None


# ── PDF ────────────────────────────────────────────────────────────────────────

def generate_pdf(display_name, cms, averages, manual):
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=15*mm, rightMargin=15*mm, topMargin=10*mm, bottomMargin=15*mm)

    brand_s  = ParagraphStyle("Brand",  fontSize=13, fontName="Helvetica-Bold",
                               textColor=colors.white, alignment=TA_CENTER, leading=18)
    title_s  = ParagraphStyle("Title",  fontSize=13, fontName="Helvetica-Bold",
                               textColor=colors.HexColor("#003366"), alignment=TA_CENTER)
    state_s  = ParagraphStyle("State",  fontSize=11, fontName="Helvetica-Bold",
                               textColor=colors.HexColor("#003366"), alignment=TA_CENTER, spaceAfter=4)
    label_s  = ParagraphStyle("Label",  fontSize=9,  fontName="Helvetica-Bold",    leading=12)
    value_s  = ParagraphStyle("Value",  fontSize=9,  fontName="Helvetica-Oblique", leading=12)
    link_s   = ParagraphStyle("Link",   fontSize=8,  fontName="Helvetica",
                               textColor=colors.HexColor("#003366"), alignment=TA_CENTER)
    footer_s = ParagraphStyle("Footer", fontSize=7,  fontName="Helvetica-Oblique",
                               textColor=colors.grey, alignment=TA_CENTER)

    story = []
    banner = Table([[Paragraph("INFINITE — Managed by MEDELITE", brand_s)]], colWidths=[180*mm])
    banner.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),colors.HexColor("#003366")),
                                 ("TOPPADDING",(0,0),(-1,-1),8),
                                 ("BOTTOMPADDING",(0,0),(-1,-1),8)]))
    story += [banner, Spacer(1,5*mm)]
    story.append(Paragraph("FACILITY ASSESSMENT SNAPSHOT", title_s))
    story.append(Paragraph(_v(cms,"state"), state_s))
    ccn = _v(cms,"cms_certification_number_ccn")
    medicare_url = f"https://www.medicare.gov/care-compare/details/nursing-home/{ccn}"
    story.append(Paragraph(f'<a href="{medicare_url}" color="#003366">View on Medicare Care Compare</a>', link_s))
    story.append(Spacer(1,3*mm))

    def row(label, value, highlight=False):
        return (Paragraph(label, label_s), Paragraph(value or "N/A", value_s), highlight)
    def star(val): return f"{val}/5" if val != "N/A" else "N/A"

    addr = ", ".join(x for x in [_v(cms,"provider_address"),_v(cms,"citytown"),
                                   _v(cms,"state"),_v(cms,"zip_code")] if x != "N/A") or "N/A"
    census = manual.get("Current Census","").strip() or _v(cms,"average_number_of_residents_per_day")

    rows = [
        row("Name of Facility",   display_name),
        row("Location",           addr),
        row("EMR",                manual.get("EMR System","") or "N/A"),
        row("Census Capacity",    _v(cms,"number_of_certified_beds")),
        row("Current Census",     census),
        row("Type of Patient",    manual.get("Type of Patient","") or "N/A"),
        row("Previous Coverage from Medelite",
            manual.get("Previous Coverage from Medelite","") or "N/A"),
        row("Previous Provider Performance from Medelite",
            manual.get("Previous Provider Performance from Medelite","") or "N/A"),
        row("Medical Coverage",   manual.get("Medical Coverage","") or "N/A"),
        row("Overall Star Rating",       star(_v(cms,"overall_rating"))),
        row("Health Inspection",         star(_v(cms,"health_inspection_rating"))),
        row("Staffing",                  star(_v(cms,"staffing_rating"))),
        row("Quality of Resident Care",  star(_v(cms,"qm_rating"))),
    ]
    for label, col, source, is_pct in HOSP_FIELDS:
        rows.append(row(label, resolve_hosp(col, source, averages, is_pct), highlight=True))
    if manual.get("Notes","").strip():
        rows.append(row("Additional Notes", manual["Notes"]))

    tbl = Table([(r[0],r[1]) for r in rows], colWidths=[80*mm,100*mm])
    style = [("BOX",(0,0),(-1,-1),0.5,colors.black),
             ("INNERGRID",(0,0),(-1,-1),0.5,colors.black),
             ("VALIGN",(0,0),(-1,-1),"TOP"),
             ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
             ("LEFTPADDING",(0,0),(-1,-1),5),("RIGHTPADDING",(0,0),(-1,-1),5)]
    for i,r in enumerate(rows):
        if r[2]: style.append(("BACKGROUND",(0,i),(-1,i),colors.HexColor("#FFFF00")))
    tbl.setStyle(TableStyle(style))
    story.append(tbl)
    story += [Spacer(1,5*mm),
              HRFlowable(width="100%",thickness=0.5,color=colors.grey),
              Paragraph("Generated by INFINITE — Managed by MEDELITE  |  Source: CMS Provider Data Catalog", footer_s)]
    doc.build(story)
    return buf.getvalue()


# ── DOCX ──────────────────────────────────────────────────────────────────────

def generate_docx(display_name, cms, averages, manual):
    """Generate Word doc using python-docx (pure Python, no Node.js required)."""
    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    # Banner
    banner_p = doc.add_paragraph()
    banner_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_p.paragraph_format.space_before = Pt(0)
    banner_p.paragraph_format.space_after  = Pt(4)
    banner_run = banner_p.add_run("INFINITE — Managed by MEDELITE")
    banner_run.bold = True
    banner_run.font.size = Pt(13)
    banner_run.font.name = 'Arial'
    banner_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    pPr = banner_p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '003366')
    pPr.append(shd)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(8)
    title_p.paragraph_format.space_after  = Pt(2)
    t_run = title_p.add_run("FACILITY ASSESSMENT SNAPSHOT")
    t_run.bold = True
    t_run.font.size = Pt(13)
    t_run.font.name = 'Arial'
    t_run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    # State
    state_val = _v(cms, "state")
    state_p = doc.add_paragraph()
    state_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    state_p.paragraph_format.space_before = Pt(0)
    state_p.paragraph_format.space_after  = Pt(6)
    s_run = state_p.add_run(state_val)
    s_run.bold = True
    s_run.font.size = Pt(11)
    s_run.font.name = 'Arial'
    s_run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    # Medicare hyperlink
    ccn = _v(cms, "cms_certification_number_ccn")
    medicare_url = f"https://www.medicare.gov/care-compare/details/nursing-home/{ccn}"
    link_p = doc.add_paragraph()
    link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    link_p.paragraph_format.space_after = Pt(8)
    _docx_add_hyperlink(link_p, medicare_url, "View on Medicare Care Compare")

    # Table
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def star(val): return f"{val}/5" if val != "N/A" else "N/A"
    addr = ", ".join(x for x in [_v(cms,"provider_address"), _v(cms,"citytown"),
                                   _v(cms,"state"), _v(cms,"zip_code")] if x != "N/A") or "N/A"
    census = manual.get("Current Census","").strip() or _v(cms,"average_number_of_residents_per_day")

    main_rows = [
        ("Name of Facility",   display_name),
        ("Location",           addr),
        ("EMR",                manual.get("EMR System","") or "N/A"),
        ("Census Capacity",    _v(cms,"number_of_certified_beds")),
        ("Current Census",     census),
        ("Type of Patient",    manual.get("Type of Patient","") or "N/A"),
        ("Previous Coverage from Medelite",
                               manual.get("Previous Coverage from Medelite","") or "N/A"),
        ("Previous Provider Performance from Medelite",
                               manual.get("Previous Provider Performance from Medelite","") or "N/A"),
        ("Medical Coverage",   manual.get("Medical Coverage","") or "N/A"),
        ("Overall Star Rating",       star(_v(cms,"overall_rating"))),
        ("Health Inspection",         star(_v(cms,"health_inspection_rating"))),
        ("Staffing",                  star(_v(cms,"staffing_rating"))),
        ("Quality of Resident Care",  star(_v(cms,"qm_rating"))),
    ]
    for label, value in main_rows:
        _docx_add_row(table, label, value, highlight=False)

    for label, col, source, is_pct in HOSP_FIELDS:
        value = resolve_hosp(col, source, averages, is_pct)
        _docx_add_row(table, label, value, highlight=True)

    if manual.get("Notes","").strip():
        _docx_add_row(table, "Additional Notes", manual["Notes"])

    # Footer
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(8)
    f_run = footer_p.add_run(
        "Generated by INFINITE — Managed by MEDELITE  |  Source: CMS Provider Data Catalog")
    f_run.italic = True
    f_run.font.size = Pt(7)
    f_run.font.name = 'Arial'
    f_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── UI helpers ─────────────────────────────────────────────────────────────────

def star_badge(val):
    try:
        n = int(float(val))
        cls = f"star-{n}"
        stars = "★" * n + "☆" * (5 - n)
        return f'<span class="{cls}">{stars}</span> {n}/5'
    except Exception:
        return "N/A"

def render_metric_cards(cms):
    ratings = [
        ("Overall Rating",         _v(cms,"overall_rating")),
        ("Health Inspection",      _v(cms,"health_inspection_rating")),
        ("Staffing",               _v(cms,"staffing_rating")),
        ("Quality of Resident Care", _v(cms,"qm_rating")),
    ]
    cols = st.columns(4)
    for col, (label, val) in zip(cols, ratings):
        try:
            n = int(float(val))
            stars = "★" * n + "☆" * (5 - n)
            color = ["","#ef4444","#f97316","#eab308","#22c55e","#16a34a"][n]
        except Exception:
            stars, color, n = "N/A", "#94a3b8", "–"
        with col:
            st.markdown(f"""
            <div class="metric-card">
              <div class="metric-label">{label}</div>
              <div class="metric-value" style="color:{color}">{stars}</div>
              <div class="metric-sub">{n} / 5</div>
            </div>""", unsafe_allow_html=True)

def render_hosp_table(averages):
    rows = []
    for label, col, source, is_pct in HOSP_FIELDS:
        val = resolve_hosp(col, source, averages, is_pct)
        rows.append({"Metric": label, "Value": val})
    df = pd.DataFrame(rows)
    # Style facility rows differently
    def highlight(row):
        styles = []
        for _ in row:
            if row["Value"] == "Not Reported":
                styles.append("color:#94a3b8;font-style:italic")
            else:
                styles.append("")
        return styles
    st.dataframe(df.style.apply(highlight, axis=1),
                 use_container_width=True, hide_index=True)


# ── Session state ──────────────────────────────────────────────────────────────
for k, v in [("cms_data",None),("averages",{"national":{},"state":{}}),
              ("facility_name",""),("display_name",""),("last_error","")]:
    if k not in st.session_state:
        st.session_state[k] = v

# ── Layout ─────────────────────────────────────────────────────────────────────
st.title("Facility Assessment Snapshot")
st.markdown("Enter a 6-digit CCN to retrieve CMS provider data and generate a report.")

col1, col2 = st.columns([1, 2])

with col1:
    ccn_raw       = st.text_input("CCN (Provider Number)", placeholder="e.g., 686123",
                                   max_chars=6)
    override_name = st.text_input("Facility Name Override (Optional)",
                                   placeholder="Leave blank to use CMS name")

    st.markdown('<div class="section-header">Manual Inputs</div>', unsafe_allow_html=True)
    manual_data = {}
    manual_data["EMR System"] = st.text_input("EMR System",
        placeholder="e.g., PointClickCare, MatrixCare")
    manual_data["Current Census"] = st.text_input("Current Census",
        placeholder="e.g., 112  (auto-filled from CMS if blank)")
    manual_data["Type of Patient"] = st.text_input("Type of Patient",
        placeholder="e.g., Long-term & Short-term")
    prev = st.selectbox("Previous Coverage from Medelite", ["Select...","Yes","No"])
    manual_data["Previous Coverage from Medelite"] = "" if prev == "Select..." else prev
    manual_data["Previous Provider Performance from Medelite"] = st.text_input(
        "Previous Provider Performance from Medelite",
        placeholder="e.g., About 30 patients/day")
    manual_data["Medical Coverage"] = st.text_input("Medical Coverage",
        placeholder="e.g., Optometry, PCP, Podiatry")
    manual_data["Notes"] = st.text_area("Additional Notes", placeholder="Any notes...")

    fetch_clicked = st.button("Fetch CMS Data", type="primary", use_container_width=True)

# ── Fetch with validation ──────────────────────────────────────────────────────
if fetch_clicked:
    ccn, err = validate_ccn(ccn_raw)
    if err:
        st.session_state.last_error = err
        st.session_state.cms_data   = None
    else:
        st.session_state.last_error = ""
        with st.spinner("Fetching provider info from CMS…"):
            cms_data = fetch_provider_data(ccn)
        if cms_data:
            st.session_state.cms_data      = cms_data
            st.session_state.facility_name = cms_data.get("provider_name","Unknown")
            st.session_state.display_name  = (
                override_name.strip() if override_name.strip()
                else st.session_state.facility_name)
            state = cms_data.get("state","")
            with st.spinner("Fetching state & national averages…"):
                st.session_state.averages = fetch_averages(state)
            st.success("✅ CMS data retrieved successfully!")
        else:
            st.session_state.cms_data = None
            st.session_state.averages = {"national":{},"state":{}}
            st.session_state.last_error = (
                f"No facility found for CCN **{ccn}**. "
                "Please double-check the number on "
                "[Medicare Care Compare](https://www.medicare.gov/care-compare/).")

if st.session_state.last_error:
    st.error(st.session_state.last_error)

# ── Results ────────────────────────────────────────────────────────────────────
if st.session_state.cms_data:
    cms          = st.session_state.cms_data
    averages     = st.session_state.averages
    facility_name = st.session_state.facility_name
    display_name  = override_name.strip() if override_name.strip() else facility_name

    with col2:
        ccn_val    = _v(cms,"cms_certification_number_ccn")
        state_code = _v(cms,"state")

        st.subheader(f"{display_name}")
        st.caption(f"CCN: {ccn_val}  ·  State: {state_code}  ·  "
                   f"{_v(cms,'provider_type')}  ·  {_v(cms,'ownership_type')}")

        addr = ", ".join(x for x in [_v(cms,"provider_address"),_v(cms,"citytown"),
                                      _v(cms,"state"),_v(cms,"zip_code")] if x != "N/A")
        st.markdown(f"📍 {addr}  ·  📞 {_v(cms,'telephone_number')}")
        st.markdown(f"🛏 **{_v(cms,'number_of_certified_beds')}** certified beds  ·  "
                    f"👥 **{_v(cms,'average_number_of_residents_per_day')}** avg residents/day")

        medicare_url = f"https://www.medicare.gov/care-compare/details/nursing-home/{ccn_val}"
        st.markdown(f"[🔗 View on Medicare Care Compare]({medicare_url})")

        st.markdown("---")
        st.markdown('<div class="section-header">Star Ratings</div>', unsafe_allow_html=True)
        render_metric_cards(cms)

        st.markdown("---")
        st.markdown('<div class="section-header">Hospitalization & ED Metrics</div>',
                    unsafe_allow_html=True)
        st.caption("Facility-specific scores are not available via the public CMS API — "
                   "shown as 'Not Reported'. State and national averages are from the "
                   "CMS State/US Averages dataset.")
        render_hosp_table(averages)

        st.markdown("---")
        st.markdown('<div class="section-header">Export Report</div>', unsafe_allow_html=True)

        dl_col1, dl_col2 = st.columns(2)

        with dl_col1:
            pdf_bytes = generate_pdf(display_name, cms, averages, manual_data)
            safe = display_name.replace(" ","_").replace("/","-")
            st.download_button(
                label="⬇ Download PDF",
                data=pdf_bytes,
                file_name=f"Snapshot_{safe}.pdf",
                mime="application/pdf",
                use_container_width=True,
                type="primary"
            )

        with dl_col2:
            try:
                docx_bytes = generate_docx(display_name, cms, averages, manual_data)
                st.download_button(
                    label="⬇ Download Word Doc (.docx)",
                    data=docx_bytes,
                    file_name=f"Snapshot_{safe}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"Word export error: {e}")

st.markdown("---")
st.markdown(
    "Powered by [CMS Provider Data Catalog](https://data.cms.gov/provider-data/dataset/4pq5-n9py) "
    "| INFINITE — Managed by MEDELITE"
)
